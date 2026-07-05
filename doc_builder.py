"""
doc_builder.py
--------------
Turns the agent's plan + drafted (and possibly self-corrected) sections into
a polished .docx using python-docx. Kept separate from agent.py so the
"how do we render a document" concern never leaks into "how do we plan and
reason" - either can change independently.
"""
from __future__ import annotations

import datetime as dt
from typing import Iterable

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_title_block(document: Document, title: str, document_type: str) -> None:
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = ACCENT

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"{document_type.title()}  |  Generated {dt.date.today().isoformat()}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # simple rule under the title block (paragraph bottom border, not a table)
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "1F3A5F",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_assumptions_box(document: Document, assumptions: Iterable[str]) -> None:
    assumptions = list(assumptions)
    if not assumptions:
        return
    document.add_heading("Assumptions Made by the Agent", level=2)
    intro = document.add_paragraph(
        "The request was ambiguous or incomplete on some points. To proceed "
        "autonomously, the agent made the following explicit assumptions:"
    )
    intro.runs[0].italic = True
    for item in assumptions:
        document.add_paragraph(item, style="List Bullet")


def _add_section(document: Document, heading: str, content: str, revised: bool) -> None:
    h = document.add_heading(heading, level=1)
    for run in h.runs:
        run.font.color.rgb = ACCENT

    if revised:
        note = document.add_paragraph()
        note_run = note.add_run("(self-corrected by agent QA pass)")
        note_run.italic = True
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    for block in content.split("\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("- ") or block.startswith("* "):
            document.add_paragraph(block[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(block)


def _add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = "Generated autonomously by AutonomousDocAgent"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def build_document(
    title: str,
    document_type: str,
    assumptions: list,
    sections: list,  # list[SectionResult]
    output_path: str,
) -> str:
    document = Document()

    doc_section = document.sections[0]
    doc_section.page_width = Inches(8.5)
    doc_section.page_height = Inches(11)
    doc_section.left_margin = Inches(1)
    doc_section.right_margin = Inches(1)

    _set_base_style(document)
    _add_title_block(document, title, document_type)
    _add_assumptions_box(document, assumptions)

    for result in sections:
        _add_section(document, result.title, result.content, result.revised)

    _add_footer(document)
    document.save(output_path)
    return output_path
